
import streamlit as st
from docx import Document
import io
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def extract_pages_by_manual_pagebreaks(doc):
    pages = []
    current_page = []
    for para in doc.paragraphs:
        current_page.append(para)
        if para._element.xpath('.//w:br[@w:type="page"]'):
            pages.append(current_page)
            current_page = []
    if current_page:
        pages.append(current_page)
    return pages

def search_keyword_in_pages(pages, keyword):
    matched = []
    for page in pages:
        full_text = "\n".join([p.text for p in page])
        if keyword in full_text:
            matched.append(page)
    return matched

def create_new_doc(pages):
    new_doc = Document()
    for page in pages:
        for para in page:
            new_para = new_doc.add_paragraph(para.text)
        run = new_doc.add_paragraph().add_run()
        run.add_break(break_type=1)
    output = io.BytesIO()
    new_doc.save(output)
    output.seek(0)
    return output

st.title("🔍 استخراج الصفحات التي تحتوي على كلمة من ملف Word")

uploaded_file = st.file_uploader("📤 ارفع ملف Word", type=["docx"])
keyword = st.text_input("🔎 أدخل الكلمة التي تبحث عنها")

if uploaded_file and keyword:
    doc = Document(uploaded_file)
    pages = extract_pages_by_manual_pagebreaks(doc)
    result_pages = search_keyword_in_pages(pages, keyword)

    if result_pages:
        output_doc = create_new_doc(result_pages)
        st.success(f"تم العثور على {len(result_pages)} صفحة تحتوي على '{keyword}'.")

        st.download_button(
            label="📥 تحميل الصفحات إلى Word",
            data=output_doc,
            file_name=f"نتائج_{keyword}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("لم يتم العثور على الكلمة في أي صفحة.")
