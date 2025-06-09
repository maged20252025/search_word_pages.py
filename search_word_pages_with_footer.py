
import streamlit as st
from docx import Document
from docx.enum.text import WD_BREAK
import io

def extract_pages_by_manual_pagebreaks(doc):
    pages = []
    current_page = []
    for para in doc.paragraphs:
        current_page.append(para)
        if para._element.xpath('.//w:br[@w:type="page"]'):
            pages.append(current_page)
            current_page = []
    if current_page:
        pages.append(current_page)
    return pages

def search_keyword_in_pages(pages, keyword):
    matched = []
    for page in pages:
        full_text = "\n".join([p.text for p in page])
        if keyword in full_text:
            matched.append(page)
    return matched

def get_footer_texts(doc):
    # نحصل على كل تذييلات الأقسام (عادةً ملف Word يحتوي على قسم واحد)
    footers = []
    for section in doc.sections:
        footer = section.footer
        footer_text = "\n".join([p.text for p in footer.paragraphs if p.text.strip()])
        footers.append(footer_text)
    return footers[0] if footers else ""

def create_new_doc(pages, footer_text):
    new_doc = Document()
    for page in pages:
        for para in page:
            new_para = new_doc.add_paragraph(para.text)
        if footer_text:
            new_doc.add_paragraph("")  # مسافة قبل التذييل
            new_doc.add_paragraph(footer_text).alignment = 2  # محاذاة لليمين
        run = new_doc.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)
    output = io.BytesIO()
    new_doc.save(output)
    output.seek(0)
    return output

st.title("🔍 استخراج الصفحات التي تحتوي على كلمة من ملف Word (مع التذييل)")

uploaded_file = st.file_uploader("📤 ارفع ملف Word", type=["docx"])
keyword = st.text_input("🔎 أدخل الكلمة التي تبحث عنها")

if uploaded_file and keyword:
    doc = Document(uploaded_file)
    footer_text = get_footer_texts(doc)
    pages = extract_pages_by_manual_pagebreaks(doc)
    result_pages = search_keyword_in_pages(pages, keyword)

    if result_pages:
        output_doc = create_new_doc(result_pages, footer_text)
        st.success(f"تم العثور على {len(result_pages)} صفحة تحتوي على '{keyword}'.")

        st.download_button(
            label="📥 تحميل الصفحات إلى Word (مع التذييل)",
            data=output_doc,
            file_name=f"نتائج_{keyword}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("لم يتم العثور على الكلمة في أي صفحة.")
